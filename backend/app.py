"""
AI Cohort Assistant – Flask backend
Stack : Flask · SQLite · ChromaDB · OpenAI · python-docx
Run   : cd backend && python app.py
"""

import os
import re
import json
import uuid
import sqlite3
import secrets
import smtplib
from datetime import datetime, timedelta
from pathlib import Path
from itertools import combinations
from email.message import EmailMessage
from html import escape as html_escape

from flask import Flask, request, jsonify, send_from_directory, send_file, Response
from flask_cors import CORS
from dotenv import load_dotenv
import chromadb
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

# ── Paths ─────────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).parent.parent
BACKEND_DIR  = Path(__file__).parent
DB_PATH      = BACKEND_DIR / "cohort.db"
CHROMA_PATH  = BACKEND_DIR / "chroma_data"
DOCS_OUT_DIR = BACKEND_DIR / "generated_docs"
DOCS_OUT_DIR.mkdir(exist_ok=True)

# ── App ───────────────────────────────────────────────────────────────────────
app = Flask(__name__)
CORS(app)

# ── OpenAI ────────────────────────────────────────────────────────────────────
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
openai_client  = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
APP_BASE_URL = os.getenv("APP_BASE_URL", "http://localhost:5000").rstrip("/")
SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")
SMTP_FROM = os.getenv("SMTP_FROM", "")

# ── ChromaDB (persistent on disk) ─────────────────────────────────────────────
chroma_client = chromadb.PersistentClient(path=str(CHROMA_PATH))


# ── SQLite helpers ────────────────────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def _smtp_ready():
    return bool(SMTP_HOST and SMTP_PORT and SMTP_FROM)


def _send_email(to_email: str, subject: str, body: str):
    if not _smtp_ready():
        return {"sent": False, "reason": "smtp_not_configured"}
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = SMTP_FROM
    msg["To"] = to_email
    msg.set_content(body)
    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=10) as smtp:
            smtp.starttls()
            if SMTP_USER:
                smtp.login(SMTP_USER, SMTP_PASS)
            smtp.send_message(msg)
        return {"sent": True}
    except Exception as e:
        return {"sent": False, "reason": str(e)[:200]}


def _mutation_not_found(entity: str, entity_id):
    return jsonify({"error": f"{entity} {entity_id} not found"}), 404


def init_db():
    with get_db() as conn:
        conn.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            username   TEXT UNIQUE NOT NULL,
            created_at TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS cohorts (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            name       TEXT UNIQUE NOT NULL,
            created_at TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS messages (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            cohort_id   INTEGER NOT NULL REFERENCES cohorts(id),
            user_id     INTEGER REFERENCES users(id),
            content     TEXT NOT NULL,
            sender_type TEXT NOT NULL CHECK(sender_type IN ('human','ai','system')),
            created_at  TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS documents (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            cohort_id   INTEGER NOT NULL REFERENCES cohorts(id),
            filename    TEXT NOT NULL,
            chunk_count INTEGER,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS generated_docs (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            cohort_id   INTEGER NOT NULL REFERENCES cohorts(id),
            user_id     INTEGER REFERENCES users(id),
            prompt      TEXT NOT NULL,
            filename    TEXT NOT NULL,
            filepath    TEXT NOT NULL,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        /* ── Tier 2: Budget Tracker ──────────────────────────────── */
        CREATE TABLE IF NOT EXISTS budgets (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id      TEXT NOT NULL,
            category        TEXT NOT NULL,
            allocated_amount REAL DEFAULT 0,
            spent_amount    REAL DEFAULT 0,
            created_at      TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS expenses (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            budget_id   INTEGER NOT NULL REFERENCES budgets(id),
            wedding_id  TEXT NOT NULL,
            vendor_name TEXT,
            amount      REAL NOT NULL,
            description TEXT,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        /* ── Tier 2: Timeline / Tasks ────────────────────────────── */
        CREATE TABLE IF NOT EXISTS wedding_tasks (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id  TEXT NOT NULL,
            title       TEXT NOT NULL,
            description TEXT,
            assigned_to TEXT,
            due_date    TEXT,
            status      TEXT DEFAULT 'pending' CHECK(status IN ('pending','in_progress','completed','overdue')),
            priority    TEXT DEFAULT 'medium' CHECK(priority IN ('low','medium','high','urgent')),
            category    TEXT,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        /* ── Tier 2: Events (for AI Coordination) ───────────────── */
        CREATE TABLE IF NOT EXISTS wedding_events (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id  TEXT NOT NULL,
            name        TEXT NOT NULL,
            start_time  TEXT,
            end_time    TEXT,
            vendor_name TEXT,
            location    TEXT,
            event_type  TEXT,
            notes       TEXT,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        /* ── Tier 2: Vendor Marketplace ──────────────────────────── */
        CREATE TABLE IF NOT EXISTS marketplace_vendors (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            name          TEXT NOT NULL,
            category      TEXT NOT NULL,
            location      TEXT,
            price_min     REAL,
            price_max     REAL,
            description   TEXT,
            rating        REAL DEFAULT 0,
            review_count  INTEGER DEFAULT 0,
            phone         TEXT,
            website       TEXT,
            image_url     TEXT,
            featured      INTEGER DEFAULT 0,
            created_at    TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS marketplace_reviews (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            vendor_id   INTEGER NOT NULL REFERENCES marketplace_vendors(id),
            wedding_id  TEXT,
            author      TEXT,
            rating      INTEGER NOT NULL CHECK(rating BETWEEN 1 AND 5),
            review_text TEXT,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS bookings (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            vendor_id       INTEGER NOT NULL REFERENCES marketplace_vendors(id),
            wedding_id      TEXT NOT NULL,
            status          TEXT DEFAULT 'pending' CHECK(status IN ('pending','confirmed','completed','cancelled')),
            amount          REAL,
            commission_fee  REAL,
            created_at      TEXT DEFAULT (datetime('now'))
        );

        /* ── Tier 3: Guest List & RSVP ───────────────────────────── */
        CREATE TABLE IF NOT EXISTS guests (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id      TEXT NOT NULL,
            name            TEXT NOT NULL,
            email           TEXT,
            phone           TEXT,
            rsvp_status     TEXT DEFAULT 'pending' CHECK(rsvp_status IN ('pending','attending','declined','maybe')),
            meal_preference TEXT,
            plus_one        INTEGER DEFAULT 0,
            table_number    INTEGER,
            group_name      TEXT,
            notes           TEXT,
            created_at      TEXT DEFAULT (datetime('now'))
        );

        /* ── Tier 3: Notifications ───────────────────────────────── */
        CREATE TABLE IF NOT EXISTS notifications (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id  TEXT NOT NULL,
            user_id     TEXT,
            type        TEXT NOT NULL,
            title       TEXT NOT NULL,
            message     TEXT,
            read_status INTEGER DEFAULT 0,
            action_url  TEXT,
            action_target TEXT,
            created_at  TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS ai_action_logs (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id   TEXT NOT NULL,
            user_id      TEXT,
            context      TEXT,
            instruction  TEXT,
            plan_json    TEXT,
            result_json  TEXT,
            status       TEXT,
            created_at   TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS guest_invite_tokens (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id   TEXT NOT NULL,
            guest_id     INTEGER NOT NULL REFERENCES guests(id),
            token        TEXT UNIQUE NOT NULL,
            sent_to      TEXT,
            expires_at   TEXT,
            status       TEXT DEFAULT 'active' CHECK(status IN ('active','used','expired','revoked')),
            last_sent_at TEXT,
            created_at   TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS guest_checkins (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id      TEXT NOT NULL,
            guest_id        INTEGER NOT NULL REFERENCES guests(id),
            token           TEXT,
            submitted_json  TEXT NOT NULL,
            created_at      TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS announcement_channels (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id   TEXT NOT NULL,
            name         TEXT NOT NULL,
            is_default   INTEGER DEFAULT 0,
            created_at   TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS announcement_messages (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            channel_id    INTEGER NOT NULL REFERENCES announcement_channels(id),
            wedding_id    TEXT NOT NULL,
            author_type   TEXT NOT NULL CHECK(author_type IN ('user','ai','system')),
            author_name   TEXT,
            message       TEXT NOT NULL,
            metadata_json TEXT,
            created_at    TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS wedding_sites (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id    TEXT NOT NULL,
            slug          TEXT UNIQUE,
            title         TEXT,
            theme         TEXT,
            content_json  TEXT,
            status        TEXT DEFAULT 'draft' CHECK(status IN ('draft','published','archived')),
            created_at    TEXT DEFAULT (datetime('now')),
            updated_at    TEXT DEFAULT (datetime('now')),
            published_at  TEXT
        );

        CREATE TABLE IF NOT EXISTS visual_plan_items (
            id                   INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id           TEXT NOT NULL,
            item_type            TEXT NOT NULL CHECK(item_type IN ('seat','stay')),
            label                TEXT NOT NULL,
            lat                  REAL,
            lng                  REAL,
            table_number         INTEGER,
            capacity             INTEGER,
            assigned_guests_json TEXT,
            details_json         TEXT,
            created_at           TEXT DEFAULT (datetime('now')),
            updated_at           TEXT DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS invitation_cards (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            wedding_id    TEXT NOT NULL,
            title         TEXT NOT NULL,
            prompt        TEXT,
            theme         TEXT,
            content_json  TEXT,
            created_by    TEXT,
            created_at    TEXT DEFAULT (datetime('now'))
        );
        """)

        # Backward-compat migration for older DBs missing action_target
        try:
            conn.execute("ALTER TABLE notifications ADD COLUMN action_target TEXT")
        except Exception:
            pass

        conn.execute("CREATE INDEX IF NOT EXISTS idx_budgets_wedding ON budgets(wedding_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_tasks_wedding ON wedding_tasks(wedding_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_guests_wedding ON guests(wedding_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_notifs_wedding ON notifications(wedding_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_guest_invites_wedding ON guest_invite_tokens(wedding_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_ann_messages_wedding ON announcement_messages(wedding_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_visual_items_wedding ON visual_plan_items(wedding_id)")

        # Seed marketplace vendors if empty
        count = conn.execute("SELECT COUNT(*) FROM marketplace_vendors").fetchone()[0]
        if count == 0:
            _seed_marketplace(conn)


def _seed_marketplace(conn):
    """Seed the marketplace with demo vendor data."""
    vendors = [
        ("Golden Hour Studio", "photographer", "Los Angeles, CA", 2500, 8000, "Award-winning wedding photography studio specializing in romantic, natural light portraits. 10+ years of experience capturing love stories across Southern California.", 4.9, 127, "(310) 555-0142", "https://goldenhour.example.com", None, 1),
        ("Petal & Vine", "florist", "Los Angeles, CA", 1500, 6000, "Artisan floral design studio creating breathtaking arrangements from locally-sourced blooms. Known for lush garden-style and modern minimalist designs.", 4.8, 89, "(310) 555-0198", "https://petalvine.example.com", None, 1),
        ("Blue Note Events", "dj-music", "Los Angeles, CA", 1200, 4500, "Premier DJ and live music entertainment. From elegant cocktail hours to packed dance floors, we curate the perfect soundtrack for your celebration.", 4.7, 156, "(323) 555-0167", "https://bluenote.example.com", None, 1),
        ("Hearth Table", "caterer", "Los Angeles, CA", 3000, 15000, "Farm-to-table catering with globally-inspired menus. Custom tasting experiences and full-service event dining for unforgettable wedding feasts.", 4.8, 203, "(310) 555-0234", "https://hearthtable.example.com", None, 1),
        ("Rosewood Garden Estate", "venue", "Pasadena, CA", 8000, 25000, "A stunning 5-acre garden estate featuring manicured lawns, a grand ballroom, and a charming courtyard perfect for ceremonies of all sizes.", 4.9, 312, "(626) 555-0101", "https://rosewood.example.com", None, 1),
        ("Sugar & Bloom Bakery", "bakery", "Santa Monica, CA", 500, 3000, "Custom wedding cakes and dessert tables that are as beautiful as they are delicious. Specializing in buttercream, fondant, and naked cake designs.", 4.6, 78, "(310) 555-0345", "https://sugarbloom.example.com", None, 0),
        ("Luminous Films", "videographer", "Los Angeles, CA", 3000, 10000, "Cinematic wedding videography that tells your story. Documentary and editorial styles with same-day highlight edits available.", 4.8, 95, "(323) 555-0289", "https://luminousfilms.example.com", None, 0),
        ("Belle & Blush", "makeup-hair", "Beverly Hills, CA", 400, 2000, "Luxury bridal beauty team offering hair styling, makeup artistry, and on-location services. Airbrush and traditional techniques available.", 4.7, 145, "(310) 555-0456", "https://belleblush.example.com", None, 0),
        ("Ever After Events", "planner", "Los Angeles, CA", 3500, 12000, "Full-service wedding planning and day-of coordination. We handle every detail so you can enjoy the journey to your big day.", 4.9, 67, "(213) 555-0178", "https://everafter.example.com", None, 0),
        ("Dreamscape Decor", "decor-rentals", "Los Angeles, CA", 1000, 8000, "Premium event rentals and custom decor installations. Furniture, lighting, linens, and styling to transform any venue into your vision.", 4.5, 112, "(323) 555-0390", "https://dreamscape.example.com", None, 0),
        ("Snap & Joy Photography", "photographer", "San Diego, CA", 1800, 5500, "Fun, candid wedding photography with a photojournalistic approach. We capture the real moments that make your day special.", 4.6, 88, "(619) 555-0201", None, None, 0),
        ("Enchanted Strings", "dj-music", "Orange County, CA", 2000, 6000, "Live string quartet and ensemble for ceremonies and receptions. Classical, modern pop covers, and custom arrangements available.", 4.9, 42, "(949) 555-0312", "https://enchanted.example.com", None, 0),
    ]
    for v in vendors:
        conn.execute("""
            INSERT INTO marketplace_vendors (name, category, location, price_min, price_max, description, rating, review_count, phone, website, image_url, featured)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
        """, v)


# ── Text / embedding helpers ──────────────────────────────────────────────────
def chunk_text(text: str, max_chars: int = 400) -> list:
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) < max_chars:
            current += ("\n\n" if current else "") + para
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)
    return chunks or [text]


def embed(texts: list) -> list:
    if not openai_client:
        raise RuntimeError("OPENAI_API_KEY not configured – check backend/.env")
    resp = openai_client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [item.embedding for item in resp.data]


def get_or_create_collection(cohort_id):
    name = f"cohort_{cohort_id}"
    try:
        return chroma_client.get_collection(name)
    except Exception:
        return chroma_client.create_collection(name)


# ── Static file serving ───────────────────────────────────────────────────────
@app.route("/")
def serve_index():
    return send_from_directory(PROJECT_ROOT, "index.html")


@app.route("/<path:filename>")
def serve_static(filename):
    return send_from_directory(PROJECT_ROOT, filename)


# ── API: health check ─────────────────────────────────────────────────────────
@app.route("/api/status")
def status():
    return jsonify({
        "status":           "ok",
        "openai_ready":     bool(OPENAI_API_KEY),
        "smtp_ready":       _smtp_ready(),
        "sqlite_db":        str(DB_PATH),
        "chroma_path":      str(CHROMA_PATH),
    })


# ── API: join / auth ──────────────────────────────────────────────────────────
@app.route("/api/join", methods=["POST"])
def join():
    data       = request.get_json() or {}
    username   = (data.get("username") or "").strip()
    cohort_name = (data.get("cohort") or "").strip()

    if not username or not cohort_name:
        return jsonify({"error": "username and cohort are required"}), 400

    with get_db() as conn:
        conn.execute("INSERT OR IGNORE INTO users (username) VALUES (?)", (username,))
        user = conn.execute("SELECT * FROM users WHERE username=?", (username,)).fetchone()

        conn.execute("INSERT OR IGNORE INTO cohorts (name) VALUES (?)", (cohort_name,))
        cohort = conn.execute("SELECT * FROM cohorts WHERE name=?", (cohort_name,)).fetchone()

    return jsonify({"user_id": user["id"], "cohort_id": cohort["id"]})


# ── API: messages ─────────────────────────────────────────────────────────────
@app.route("/api/messages")
def get_messages():
    cohort_id = request.args.get("cohort_id")
    with get_db() as conn:
        rows = conn.execute("""
            SELECT m.id, m.content, m.sender_type, m.created_at, u.username
            FROM   messages m
            LEFT JOIN users u ON m.user_id = u.id
            WHERE  m.cohort_id = ?
            ORDER  BY m.created_at
        """, (cohort_id,)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/message", methods=["POST"])
def send_message():
    data = request.get_json() or {}
    with get_db() as conn:
        conn.execute(
            "INSERT INTO messages (cohort_id, user_id, content, sender_type) VALUES (?,?,?,?)",
            (data.get("cohort_id"), data.get("user_id"), data.get("content"), data.get("sender_type", "human"))
        )
    return jsonify({"success": True})


# ── API: document upload → chunk → embed → ChromaDB ──────────────────────────
@app.route("/api/upload", methods=["POST"])
def upload():
    cohort_id = request.form.get("cohort_id")
    file      = request.files.get("file")

    if not file:
        return jsonify({"error": "no file provided"}), 400
    if not cohort_id:
        return jsonify({"error": "cohort_id required"}), 400

    text     = file.read().decode("utf-8", errors="replace")
    filename = file.filename
    chunks   = chunk_text(text)

    embeddings = embed(chunks)
    collection = get_or_create_collection(cohort_id)
    ids = [f"{filename}_{uuid.uuid4().hex[:8]}_{i}" for i in range(len(chunks))]
    collection.add(
        documents=chunks,
        embeddings=embeddings,
        ids=ids,
        metadatas=[{"filename": filename, "chunk_idx": i} for i in range(len(chunks))]
    )

    with get_db() as conn:
        conn.execute(
            "INSERT INTO documents (cohort_id, filename, chunk_count) VALUES (?,?,?)",
            (cohort_id, filename, len(chunks))
        )

    return jsonify({"success": True, "filename": filename, "chunks": len(chunks)})


# ══════════════════════════════════════════════════════════════════════════════
# BUDGET TRACKER API
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/budgets")
def get_budgets():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM budgets WHERE wedding_id=? ORDER BY category", (wedding_id,)
        ).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/budgets", methods=["POST"])
def create_budget():
    data = request.get_json() or {}
    wedding_id = data.get("wedding_id")
    category = (data.get("category") or "").strip()
    allocated = data.get("allocated_amount", 0)
    if not wedding_id or not category:
        return jsonify({"error": "wedding_id and category required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO budgets (wedding_id, category, allocated_amount) VALUES (?,?,?)",
            (wedding_id, category, allocated)
        )
        bid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": bid})


@app.route("/api/budgets/<int:budget_id>", methods=["PUT"])
def update_budget(budget_id):
    data = request.get_json() or {}
    if "allocated_amount" not in data:
        return jsonify({"error": "allocated_amount is required"}), 400
    with get_db() as conn:
        cur = conn.execute(
            "UPDATE budgets SET allocated_amount=? WHERE id=?",
            (data.get("allocated_amount", 0), budget_id)
        )
    if cur.rowcount == 0:
        return _mutation_not_found("budget", budget_id)
    return jsonify({"success": True, "updated": cur.rowcount})


@app.route("/api/budgets/<int:budget_id>", methods=["DELETE"])
def delete_budget(budget_id):
    with get_db() as conn:
        conn.execute("DELETE FROM expenses WHERE budget_id=?", (budget_id,))
        cur = conn.execute("DELETE FROM budgets WHERE id=?", (budget_id,))
    if cur.rowcount == 0:
        return _mutation_not_found("budget", budget_id)
    return jsonify({"success": True, "deleted": cur.rowcount})


@app.route("/api/expenses")
def get_expenses():
    wedding_id = request.args.get("wedding_id")
    budget_id = request.args.get("budget_id")
    with get_db() as conn:
        if budget_id:
            rows = conn.execute("SELECT * FROM expenses WHERE budget_id=? ORDER BY created_at DESC", (budget_id,)).fetchall()
        elif wedding_id:
            rows = conn.execute("SELECT * FROM expenses WHERE wedding_id=? ORDER BY created_at DESC", (wedding_id,)).fetchall()
        else:
            return jsonify({"error": "wedding_id or budget_id required"}), 400
    return jsonify([dict(r) for r in rows])


@app.route("/api/expenses", methods=["POST"])
def create_expense():
    data = request.get_json() or {}
    budget_id = data.get("budget_id")
    wedding_id = data.get("wedding_id")
    amount = data.get("amount", 0)
    if not budget_id or not wedding_id:
        return jsonify({"error": "budget_id and wedding_id required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO expenses (budget_id, wedding_id, vendor_name, amount, description) VALUES (?,?,?,?,?)",
            (budget_id, wedding_id, data.get("vendor_name", ""), amount, data.get("description", ""))
        )
        # Update spent_amount on budget
        total = conn.execute("SELECT COALESCE(SUM(amount),0) FROM expenses WHERE budget_id=?", (budget_id,)).fetchone()[0]
        conn.execute("UPDATE budgets SET spent_amount=? WHERE id=?", (total, budget_id))
    return jsonify({"success": True})


@app.route("/api/expenses/<int:expense_id>", methods=["DELETE"])
def delete_expense(expense_id):
    with get_db() as conn:
        row = conn.execute("SELECT budget_id FROM expenses WHERE id=?", (expense_id,)).fetchone()
        if not row:
            return _mutation_not_found("expense", expense_id)
        conn.execute("DELETE FROM expenses WHERE id=?", (expense_id,))
        total = conn.execute("SELECT COALESCE(SUM(amount),0) FROM expenses WHERE budget_id=?", (row["budget_id"],)).fetchone()[0]
        conn.execute("UPDATE budgets SET spent_amount=? WHERE id=?", (total, row["budget_id"]))
    return jsonify({"success": True, "deleted": 1})


@app.route("/api/budget-summary")
def budget_summary():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM budgets WHERE wedding_id=?", (wedding_id,)).fetchall()
        total_allocated = sum(r["allocated_amount"] for r in rows)
        total_spent = sum(r["spent_amount"] for r in rows)
    return jsonify({
        "total_allocated": total_allocated,
        "total_spent": total_spent,
        "remaining": total_allocated - total_spent,
        "categories": [dict(r) for r in rows]
    })


# ══════════════════════════════════════════════════════════════════════════════
# TIMELINE / TASKS API
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/tasks")
def get_tasks():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM wedding_tasks WHERE wedding_id=? ORDER BY due_date, priority DESC", (wedding_id,)
        ).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/tasks", methods=["POST"])
def create_task():
    data = request.get_json() or {}
    wedding_id = data.get("wedding_id")
    title = (data.get("title") or "").strip()
    if not wedding_id or not title:
        return jsonify({"error": "wedding_id and title required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO wedding_tasks (wedding_id, title, description, assigned_to, due_date, status, priority, category) VALUES (?,?,?,?,?,?,?,?)",
            (wedding_id, title, data.get("description", ""), data.get("assigned_to", ""),
             data.get("due_date", ""), data.get("status", "pending"),
             data.get("priority", "medium"), data.get("category", ""))
        )
        tid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": tid})


@app.route("/api/tasks/<int:task_id>", methods=["PUT"])
def update_task(task_id):
    data = request.get_json() or {}
    fields, vals = [], []
    for f in ["title", "description", "assigned_to", "due_date", "status", "priority", "category"]:
        if f in data:
            fields.append(f"{f}=?")
            vals.append(data[f])
    if not fields:
        return jsonify({"error": "no fields to update"}), 400
    vals.append(task_id)
    with get_db() as conn:
        cur = conn.execute(f"UPDATE wedding_tasks SET {','.join(fields)} WHERE id=?", vals)
    if cur.rowcount == 0:
        return _mutation_not_found("task", task_id)
    return jsonify({"success": True, "updated": cur.rowcount})


@app.route("/api/tasks/<int:task_id>", methods=["DELETE"])
def delete_task(task_id):
    with get_db() as conn:
        cur = conn.execute("DELETE FROM wedding_tasks WHERE id=?", (task_id,))
    if cur.rowcount == 0:
        return _mutation_not_found("task", task_id)
    return jsonify({"success": True, "deleted": cur.rowcount})


@app.route("/api/task-summary")
def task_summary():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        total = conn.execute("SELECT COUNT(*) FROM wedding_tasks WHERE wedding_id=?", (wedding_id,)).fetchone()[0]
        completed = conn.execute("SELECT COUNT(*) FROM wedding_tasks WHERE wedding_id=? AND status='completed'", (wedding_id,)).fetchone()[0]
        in_progress = conn.execute("SELECT COUNT(*) FROM wedding_tasks WHERE wedding_id=? AND status='in_progress'", (wedding_id,)).fetchone()[0]
        overdue = conn.execute("SELECT COUNT(*) FROM wedding_tasks WHERE wedding_id=? AND status!='completed' AND due_date < date('now')", (wedding_id,)).fetchone()[0]
    return jsonify({"total": total, "completed": completed, "in_progress": in_progress, "overdue": overdue, "pending": total - completed - in_progress})


# ══════════════════════════════════════════════════════════════════════════════
# WEDDING EVENTS API (for AI Coordination)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/events")
def get_events():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM wedding_events WHERE wedding_id=? ORDER BY start_time", (wedding_id,)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/events", methods=["POST"])
def create_event():
    data = request.get_json() or {}
    wedding_id = data.get("wedding_id")
    name = (data.get("name") or "").strip()
    if not wedding_id or not name:
        return jsonify({"error": "wedding_id and name required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO wedding_events (wedding_id, name, start_time, end_time, vendor_name, location, event_type, notes) VALUES (?,?,?,?,?,?,?,?)",
            (wedding_id, name, data.get("start_time", ""), data.get("end_time", ""),
             data.get("vendor_name", ""), data.get("location", ""),
             data.get("event_type", ""), data.get("notes", ""))
        )
        eid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": eid})


@app.route("/api/events/<int:event_id>", methods=["PUT"])
def update_event(event_id):
    data = request.get_json() or {}
    fields, vals = [], []
    for f in ["name", "start_time", "end_time", "vendor_name", "location", "event_type", "notes"]:
        if f in data:
            fields.append(f"{f}=?")
            vals.append(data[f])
    if not fields:
        return jsonify({"error": "no fields to update"}), 400
    vals.append(event_id)
    with get_db() as conn:
        cur = conn.execute(f"UPDATE wedding_events SET {','.join(fields)} WHERE id=?", vals)
    if cur.rowcount == 0:
        return _mutation_not_found("event", event_id)
    return jsonify({"success": True, "updated": cur.rowcount})


@app.route("/api/events/<int:event_id>", methods=["DELETE"])
def delete_event(event_id):
    with get_db() as conn:
        cur = conn.execute("DELETE FROM wedding_events WHERE id=?", (event_id,))
    if cur.rowcount == 0:
        return _mutation_not_found("event", event_id)
    return jsonify({"success": True, "deleted": cur.rowcount})


@app.route("/api/detect-conflicts")
def detect_conflicts():
    """AI Coordination: detect scheduling conflicts in wedding events."""
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        events = conn.execute("SELECT * FROM wedding_events WHERE wedding_id=? AND start_time != '' AND end_time != '' ORDER BY start_time", (wedding_id,)).fetchall()

    conflicts = []
    event_list = [dict(e) for e in events]
    for e1, e2 in combinations(event_list, 2):
        try:
            s1, e1_end = e1["start_time"], e1["end_time"]
            s2, e2_end = e2["start_time"], e2["end_time"]
            if s1 < e2_end and s2 < e1_end:
                conflicts.append({
                    "event1": e1["name"],
                    "event2": e2["name"],
                    "event1_time": f"{s1} - {e1_end}",
                    "event2_time": f"{s2} - {e2_end}",
                    "type": "overlap",
                    "suggestion": f'"{e1["name"]}" overlaps with "{e2["name"]}". Consider adjusting times.'
                })
        except Exception:
            pass

    return jsonify({"conflicts": conflicts, "total_events": len(event_list)})


# ══════════════════════════════════════════════════════════════════════════════
# VENDOR MARKETPLACE API
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/marketplace/vendors")
def marketplace_vendors():
    category = request.args.get("category", "")
    location = request.args.get("location", "")
    price_max = request.args.get("price_max", "")
    search = request.args.get("search", "")
    sort_by = request.args.get("sort", "rating")

    query = "SELECT * FROM marketplace_vendors WHERE 1=1"
    params = []

    if category:
        query += " AND category=?"
        params.append(category)
    if location:
        query += " AND location LIKE ?"
        params.append(f"%{location}%")
    if price_max:
        try:
            query += " AND price_min <= ?"
            params.append(float(price_max))
        except ValueError:
            pass
    if search:
        query += " AND (name LIKE ? OR description LIKE ?)"
        params.extend([f"%{search}%", f"%{search}%"])

    if sort_by == "price_low":
        query += " ORDER BY price_min ASC"
    elif sort_by == "price_high":
        query += " ORDER BY price_max DESC"
    elif sort_by == "reviews":
        query += " ORDER BY review_count DESC"
    else:
        query += " ORDER BY featured DESC, rating DESC"

    with get_db() as conn:
        rows = conn.execute(query, params).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/marketplace/vendors/<int:vendor_id>")
def marketplace_vendor_detail(vendor_id):
    with get_db() as conn:
        vendor = conn.execute("SELECT * FROM marketplace_vendors WHERE id=?", (vendor_id,)).fetchone()
        if not vendor:
            return jsonify({"error": "Vendor not found"}), 404
        reviews = conn.execute(
            "SELECT * FROM marketplace_reviews WHERE vendor_id=? ORDER BY created_at DESC LIMIT 10", (vendor_id,)
        ).fetchall()
    return jsonify({"vendor": dict(vendor), "reviews": [dict(r) for r in reviews]})


@app.route("/api/marketplace/reviews", methods=["POST"])
def create_marketplace_review():
    data = request.get_json() or {}
    vendor_id = data.get("vendor_id")
    rating = data.get("rating")
    if not vendor_id or not rating:
        return jsonify({"error": "vendor_id and rating required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO marketplace_reviews (vendor_id, wedding_id, author, rating, review_text) VALUES (?,?,?,?,?)",
            (vendor_id, data.get("wedding_id", ""), data.get("author", ""), rating, data.get("review_text", ""))
        )
        # Update vendor rating
        stats = conn.execute(
            "SELECT AVG(rating) as avg_r, COUNT(*) as cnt FROM marketplace_reviews WHERE vendor_id=?", (vendor_id,)
        ).fetchone()
        conn.execute(
            "UPDATE marketplace_vendors SET rating=?, review_count=? WHERE id=?",
            (round(stats["avg_r"], 1), stats["cnt"], vendor_id)
        )
    return jsonify({"success": True})


@app.route("/api/marketplace/compare")
def compare_vendors():
    ids = request.args.get("ids", "")
    if not ids:
        return jsonify({"error": "ids parameter required (comma-separated)"}), 400
    id_list = [i.strip() for i in ids.split(",") if i.strip()]
    placeholders = ",".join(["?"] * len(id_list))
    with get_db() as conn:
        rows = conn.execute(f"SELECT * FROM marketplace_vendors WHERE id IN ({placeholders})", id_list).fetchall()
    return jsonify([dict(r) for r in rows])


# ══════════════════════════════════════════════════════════════════════════════
# GUEST LIST & RSVP API
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/guests")
def get_guests():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM guests WHERE wedding_id=? ORDER BY name", (wedding_id,)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/guests", methods=["POST"])
def create_guest():
    data = request.get_json() or {}
    wedding_id = data.get("wedding_id")
    name = (data.get("name") or "").strip()
    if not wedding_id or not name:
        return jsonify({"error": "wedding_id and name required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO guests (wedding_id, name, email, phone, rsvp_status, meal_preference, plus_one, table_number, group_name, notes) VALUES (?,?,?,?,?,?,?,?,?,?)",
            (wedding_id, name, data.get("email", ""), data.get("phone", ""),
             data.get("rsvp_status", "pending"), data.get("meal_preference", ""),
             data.get("plus_one", 0), data.get("table_number"),
             data.get("group_name", ""), data.get("notes", ""))
        )
        gid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": gid})


@app.route("/api/guests/<int:guest_id>", methods=["PUT"])
def update_guest(guest_id):
    data = request.get_json() or {}
    fields, vals = [], []
    for f in ["name", "email", "phone", "rsvp_status", "meal_preference", "plus_one", "table_number", "group_name", "notes"]:
        if f in data:
            fields.append(f"{f}=?")
            vals.append(data[f])
    if not fields:
        return jsonify({"error": "no fields to update"}), 400
    vals.append(guest_id)
    with get_db() as conn:
        cur = conn.execute(f"UPDATE guests SET {','.join(fields)} WHERE id=?", vals)
    if cur.rowcount == 0:
        return _mutation_not_found("guest", guest_id)
    return jsonify({"success": True, "updated": cur.rowcount})


@app.route("/api/guests/<int:guest_id>", methods=["DELETE"])
def delete_guest(guest_id):
    with get_db() as conn:
        cur = conn.execute("DELETE FROM guests WHERE id=?", (guest_id,))
    if cur.rowcount == 0:
        return _mutation_not_found("guest", guest_id)
    return jsonify({"success": True, "deleted": cur.rowcount})


@app.route("/api/guest-summary")
def guest_summary():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        total = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=?", (wedding_id,)).fetchone()[0]
        attending = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=? AND rsvp_status='attending'", (wedding_id,)).fetchone()[0]
        declined = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=? AND rsvp_status='declined'", (wedding_id,)).fetchone()[0]
        pending = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=? AND rsvp_status='pending'", (wedding_id,)).fetchone()[0]
        plus_ones = conn.execute("SELECT COALESCE(SUM(plus_one),0) FROM guests WHERE wedding_id=? AND rsvp_status='attending'", (wedding_id,)).fetchone()[0]
        meals = conn.execute("SELECT meal_preference, COUNT(*) as cnt FROM guests WHERE wedding_id=? AND rsvp_status='attending' AND meal_preference != '' GROUP BY meal_preference", (wedding_id,)).fetchall()
    return jsonify({
        "total": total, "attending": attending, "declined": declined, "pending": pending,
        "plus_ones": plus_ones, "total_attending": attending + plus_ones,
        "meals": {r["meal_preference"]: r["cnt"] for r in meals}
    })


# ══════════════════════════════════════════════════════════════════════════════
# NOTIFICATIONS API
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/notifications")
def get_notifications():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM notifications WHERE wedding_id=? ORDER BY created_at DESC LIMIT 50", (wedding_id,)
        ).fetchall()
        unread = conn.execute(
            "SELECT COUNT(*) FROM notifications WHERE wedding_id=? AND read_status=0", (wedding_id,)
        ).fetchone()[0]
    notifications = []
    for r in rows:
        item = dict(r)
        raw_target = item.get("action_target") or ""
        if raw_target:
            try:
                item["action_target"] = json.loads(raw_target)
            except Exception:
                item["action_target"] = raw_target
        notifications.append(item)
    return jsonify({"notifications": notifications, "unread": unread})


@app.route("/api/notifications", methods=["POST"])
def create_notification():
    data = request.get_json() or {}
    wedding_id = data.get("wedding_id")
    ntype = data.get("type", "info")
    title = (data.get("title") or "").strip()
    if not wedding_id or not title:
        return jsonify({"error": "wedding_id and title required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO notifications (wedding_id, user_id, type, title, message, action_url, action_target) VALUES (?,?,?,?,?,?,?)",
            (wedding_id, data.get("user_id", ""), ntype, title,
             data.get("message", ""), data.get("action_url", ""),
             json.dumps(data.get("action_target")) if data.get("action_target") is not None else "")
        )
    return jsonify({"success": True})


@app.route("/api/notifications/<int:notif_id>/read", methods=["PUT"])
def mark_notification_read(notif_id):
    with get_db() as conn:
        cur = conn.execute("UPDATE notifications SET read_status=1 WHERE id=?", (notif_id,))
    if cur.rowcount == 0:
        return _mutation_not_found("notification", notif_id)
    return jsonify({"success": True, "updated": cur.rowcount})


@app.route("/api/notifications/read-all", methods=["PUT"])
def mark_all_read():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        cur = conn.execute("UPDATE notifications SET read_status=1 WHERE wedding_id=?", (wedding_id,))
    return jsonify({"success": True, "updated": cur.rowcount})


def _ensure_default_announcement_channel(conn, wedding_id: str):
    row = conn.execute(
        "SELECT * FROM announcement_channels WHERE wedding_id=? AND is_default=1 ORDER BY id LIMIT 1",
        (wedding_id,),
    ).fetchone()
    if row:
        return row["id"]
    conn.execute(
        "INSERT INTO announcement_channels (wedding_id, name, is_default) VALUES (?,?,1)",
        (wedding_id, "Announcements"),
    )
    return conn.execute("SELECT last_insert_rowid()").fetchone()[0]


@app.route("/api/ai-actions/log", methods=["POST"])
def create_ai_action_log():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        conn.execute(
            """
            INSERT INTO ai_action_logs (wedding_id, user_id, context, instruction, plan_json, result_json, status)
            VALUES (?,?,?,?,?,?,?)
            """,
            (
                wedding_id,
                str(data.get("user_id", "")),
                str(data.get("context", "")),
                str(data.get("instruction", "")),
                json.dumps(data.get("plan") or {}),
                json.dumps(data.get("result") or {}),
                str(data.get("status", "")),
            ),
        )
        log_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": log_id})


@app.route("/api/ai-actions")
def list_ai_action_logs():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM ai_action_logs WHERE wedding_id=? ORDER BY created_at DESC LIMIT 200",
            (wedding_id,),
        ).fetchall()
    parsed = []
    for row in rows:
        item = dict(row)
        for field in ["plan_json", "result_json"]:
            try:
                item[field] = json.loads(item.get(field) or "{}")
            except Exception:
                pass
        parsed.append(item)
    return jsonify(parsed)


@app.route("/api/guests/import", methods=["POST"])
def import_guests():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    rows = data.get("rows") or []
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    if not isinstance(rows, list):
        return jsonify({"error": "rows must be an array"}), 400

    created, updated, skipped = 0, 0, 0
    errors = []

    with get_db() as conn:
        for idx, raw in enumerate(rows, start=1):
            row = raw if isinstance(raw, dict) else {}
            name = str(row.get("name", "")).strip()
            email = str(row.get("email", "")).strip()
            phone = str(row.get("phone", "")).strip()
            if not name:
                skipped += 1
                errors.append({"row": idx, "error": "Missing name"})
                continue

            existing = None
            if email:
                existing = conn.execute(
                    "SELECT id FROM guests WHERE wedding_id=? AND lower(email)=lower(?) LIMIT 1",
                    (wedding_id, email),
                ).fetchone()
            if not existing and phone:
                existing = conn.execute(
                    "SELECT id FROM guests WHERE wedding_id=? AND name=? AND phone=? LIMIT 1",
                    (wedding_id, name, phone),
                ).fetchone()

            payload = (
                name,
                email,
                phone,
                str(row.get("rsvp_status") or "pending"),
                str(row.get("meal_preference") or ""),
                int(row.get("plus_one") or 0),
                row.get("table_number"),
                str(row.get("group_name") or ""),
                str(row.get("notes") or ""),
            )

            if existing:
                conn.execute(
                    """
                    UPDATE guests
                    SET name=?, email=?, phone=?, rsvp_status=?, meal_preference=?, plus_one=?, table_number=?, group_name=?, notes=?
                    WHERE id=?
                    """,
                    payload + (existing["id"],),
                )
                updated += 1
            else:
                conn.execute(
                    """
                    INSERT INTO guests (wedding_id, name, email, phone, rsvp_status, meal_preference, plus_one, table_number, group_name, notes)
                    VALUES (?,?,?,?,?,?,?,?,?,?)
                    """,
                    (wedding_id,) + payload,
                )
                created += 1

    return jsonify({
        "success": True,
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "errors": errors[:50],
    })


@app.route("/api/guest-invites/send", methods=["POST"])
def send_guest_invites():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    guest_ids = data.get("guest_ids") or []
    send_email = bool(data.get("send_email", True))

    with get_db() as conn:
        query = "SELECT * FROM guests WHERE wedding_id=?"
        args = [wedding_id]
        if isinstance(guest_ids, list) and guest_ids:
            placeholders = ",".join(["?"] * len(guest_ids))
            query += f" AND id IN ({placeholders})"
            args.extend([int(gid) for gid in guest_ids])
        guests = conn.execute(query, args).fetchall()
        if not guests:
            return jsonify({"error": "No guests found for invite dispatch"}), 404

        channel_id = _ensure_default_announcement_channel(conn, wedding_id)
        results = []
        for g in guests:
            token_row = conn.execute(
                """
                SELECT * FROM guest_invite_tokens
                WHERE wedding_id=? AND guest_id=? AND status='active'
                ORDER BY created_at DESC LIMIT 1
                """,
                (wedding_id, g["id"]),
            ).fetchone()
            token = token_row["token"] if token_row else secrets.token_urlsafe(24)
            expires_at = (datetime.utcnow() + timedelta(days=30)).isoformat()
            if token_row:
                conn.execute(
                    "UPDATE guest_invite_tokens SET expires_at=?, last_sent_at=datetime('now') WHERE id=?",
                    (expires_at, token_row["id"]),
                )
            else:
                conn.execute(
                    """
                    INSERT INTO guest_invite_tokens (wedding_id, guest_id, token, sent_to, expires_at, status, last_sent_at)
                    VALUES (?,?,?,?,?,'active',datetime('now'))
                    """,
                    (wedding_id, g["id"], token, g["email"] or "", expires_at),
                )

            invite_link = f"{APP_BASE_URL}/guest/checkin/{token}"
            email_result = {"sent": False, "reason": "email_skipped"}
            if send_email and g["email"]:
                body = (
                    f"Hi {g['name']},\n\n"
                    f"You are invited! Please use your RSVP/check-in link:\n{invite_link}\n\n"
                    "This link is unique to you.\n"
                )
                email_result = _send_email(g["email"], "Wedding RSVP & Check-in Link", body)

            results.append({
                "guest_id": g["id"],
                "name": g["name"],
                "email": g["email"],
                "invite_link": invite_link,
                "email_result": email_result,
            })

        conn.execute(
            """
            INSERT INTO announcement_messages (channel_id, wedding_id, author_type, author_name, message, metadata_json)
            VALUES (?,?,?,?,?,?)
            """,
            (
                channel_id,
                wedding_id,
                "ai",
                "AI Assistant",
                f"Invitation links generated for {len(results)} guest(s).",
                json.dumps({"type": "invite_dispatch", "count": len(results)}),
            ),
        )

    return jsonify({"success": True, "results": results, "smtp_ready": _smtp_ready()})


@app.route("/api/guest-self-checkin/<token>")
def get_guest_self_checkin(token):
    with get_db() as conn:
        row = conn.execute(
            """
            SELECT git.*, g.name, g.email, g.phone, g.rsvp_status, g.meal_preference, g.plus_one, g.table_number, g.group_name, g.notes
            FROM guest_invite_tokens git
            JOIN guests g ON g.id = git.guest_id
            WHERE git.token=? AND git.status='active'
            ORDER BY git.created_at DESC LIMIT 1
            """,
            (token,),
        ).fetchone()
        if not row:
            return jsonify({"error": "Invite token not found or expired"}), 404
    item = dict(row)
    return jsonify({
        "wedding_id": item["wedding_id"],
        "guest": {
            "id": item["guest_id"],
            "name": item["name"],
            "email": item["email"],
            "phone": item["phone"],
            "rsvp_status": item["rsvp_status"],
            "meal_preference": item["meal_preference"],
            "plus_one": item["plus_one"],
            "table_number": item["table_number"],
            "group_name": item["group_name"],
            "notes": item["notes"],
        },
    })


@app.route("/api/guest-self-checkin/<token>", methods=["POST"])
def submit_guest_self_checkin(token):
    data = request.get_json() or {}
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM guest_invite_tokens WHERE token=? AND status='active' ORDER BY created_at DESC LIMIT 1",
            (token,),
        ).fetchone()
        if not row:
            return jsonify({"error": "Invite token not found or expired"}), 404
        guest_id = row["guest_id"]
        patch = {}
        for f in ["name", "email", "phone", "rsvp_status", "meal_preference", "plus_one", "table_number", "group_name", "notes"]:
            if f in data:
                patch[f] = data[f]
        if patch:
            fields = [f"{k}=?" for k in patch.keys()]
            vals = list(patch.values()) + [guest_id]
            conn.execute(f"UPDATE guests SET {','.join(fields)} WHERE id=?", vals)
        conn.execute(
            "INSERT INTO guest_checkins (wedding_id, guest_id, token, submitted_json) VALUES (?,?,?,?)",
            (row["wedding_id"], guest_id, token, json.dumps(data)),
        )
        conn.execute("UPDATE guest_invite_tokens SET status='used' WHERE id=?", (row["id"],))
    return jsonify({"success": True})


@app.route("/guest/checkin/<token>")
def guest_checkin_page(token):
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Guest Check-in</title>
  <style>
    body {{ font-family: Inter, Arial, sans-serif; margin: 0; background: #f6f2fb; color: #1f1730; }}
    .wrap {{ max-width: 620px; margin: 36px auto; background: #fff; border-radius: 14px; padding: 22px; box-shadow: 0 6px 28px rgba(63,39,99,.12); }}
    h1 {{ margin-top: 0; font-size: 1.35rem; }}
    label {{ display:block; margin-top: 10px; font-size: .85rem; font-weight: 600; }}
    input, select, textarea {{ width: 100%; margin-top: 6px; padding: 10px; border: 1px solid #d7cde8; border-radius: 10px; font: inherit; }}
    button {{ margin-top: 14px; border: 0; background: #6f49b6; color: #fff; padding: 10px 16px; border-radius: 999px; cursor: pointer; font-weight: 600; }}
    .ok {{ color: #198754; font-weight: 700; margin-top: 10px; }}
    .err {{ color: #b23b4f; margin-top: 10px; }}
  </style>
</head>
<body>
  <div class="wrap">
    <h1>Wedding Guest RSVP Check-in</h1>
    <p>Use this form to confirm your latest details.</p>
    <form id="f">
      <label>Name<input id="name" /></label>
      <label>Email<input id="email" type="email" /></label>
      <label>Phone<input id="phone" /></label>
      <label>RSVP
        <select id="rsvp_status"><option value="pending">Pending</option><option value="attending">Attending</option><option value="declined">Declined</option><option value="maybe">Maybe</option></select>
      </label>
      <label>Meal Preference<input id="meal_preference" placeholder="Vegetarian / Vegan / etc" /></label>
      <label>Plus One
        <select id="plus_one"><option value="0">No</option><option value="1">Yes</option></select>
      </label>
      <label>Table Number<input id="table_number" type="number" min="1" /></label>
      <label>Group Name<input id="group_name" /></label>
      <label>Notes<textarea id="notes" rows="3"></textarea></label>
      <button type="submit">Save Check-in</button>
    </form>
    <p id="msg"></p>
  </div>
  <script>
    const token = {json.dumps(token)};
    const msg = document.getElementById("msg");
    fetch(`/api/guest-self-checkin/${{token}}`).then(r => r.json()).then(d => {{
      if (d.error) {{ msg.className='err'; msg.textContent=d.error; return; }}
      const g = d.guest || {{}};
      ["name","email","phone","rsvp_status","meal_preference","plus_one","table_number","group_name","notes"].forEach(k => {{
        const el = document.getElementById(k); if (!el) return;
        if (g[k] !== undefined && g[k] !== null) el.value = g[k];
      }});
    }}).catch(() => {{ msg.className='err'; msg.textContent='Failed to load invite details.'; }});
    document.getElementById("f").addEventListener("submit", async (e) => {{
      e.preventDefault();
      const payload = {{}};
      ["name","email","phone","rsvp_status","meal_preference","plus_one","table_number","group_name","notes"].forEach(k => payload[k]=document.getElementById(k).value);
      payload.plus_one = parseInt(payload.plus_one || "0", 10) || 0;
      payload.table_number = payload.table_number ? parseInt(payload.table_number, 10) : null;
      const res = await fetch(`/api/guest-self-checkin/${{token}}`, {{ method: "POST", headers: {{ "Content-Type":"application/json" }}, body: JSON.stringify(payload) }});
      const out = await res.json();
      if (res.ok) {{ msg.className='ok'; msg.textContent='Saved. Thank you!'; }}
      else {{ msg.className='err'; msg.textContent = out.error || 'Save failed.'; }}
    }});
  </script>
</body>
</html>"""
    return Response(html, mimetype="text/html")


@app.route("/api/announcements/channels")
def list_announcement_channels():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        _ensure_default_announcement_channel(conn, wedding_id)
        rows = conn.execute(
            "SELECT * FROM announcement_channels WHERE wedding_id=? ORDER BY is_default DESC, id ASC",
            (wedding_id,),
        ).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/announcements/channels", methods=["POST"])
def create_announcement_channel():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    name = (data.get("name") or "").strip()
    if not wedding_id or not name:
        return jsonify({"error": "wedding_id and name required"}), 400
    with get_db() as conn:
        conn.execute(
            "INSERT INTO announcement_channels (wedding_id, name, is_default) VALUES (?,?,0)",
            (wedding_id, name),
        )
        channel_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": channel_id})


@app.route("/api/announcements/messages")
def list_announcement_messages():
    wedding_id = request.args.get("wedding_id")
    channel_id = request.args.get("channel_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        if channel_id:
            rows = conn.execute(
                "SELECT * FROM announcement_messages WHERE wedding_id=? AND channel_id=? ORDER BY created_at DESC LIMIT 200",
                (wedding_id, channel_id),
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM announcement_messages WHERE wedding_id=? ORDER BY created_at DESC LIMIT 200",
                (wedding_id,),
            ).fetchall()
    out = []
    for r in rows:
        item = dict(r)
        try:
            item["metadata_json"] = json.loads(item.get("metadata_json") or "{}")
        except Exception:
            pass
        out.append(item)
    return jsonify(out)


@app.route("/api/announcements/messages", methods=["POST"])
def create_announcement_message():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    message = (data.get("message") or "").strip()
    if not wedding_id or not message:
        return jsonify({"error": "wedding_id and message required"}), 400
    with get_db() as conn:
        channel_id = data.get("channel_id") or _ensure_default_announcement_channel(conn, wedding_id)
        conn.execute(
            """
            INSERT INTO announcement_messages (channel_id, wedding_id, author_type, author_name, message, metadata_json)
            VALUES (?,?,?,?,?,?)
            """,
            (
                channel_id,
                wedding_id,
                data.get("author_type", "user"),
                data.get("author_name", ""),
                message,
                json.dumps(data.get("metadata") or {}),
            ),
        )
        msg_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": msg_id})


@app.route("/api/wedding-sites")
def get_wedding_sites():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM wedding_sites WHERE wedding_id=? ORDER BY updated_at DESC",
            (wedding_id,),
        ).fetchall()
    out = []
    for row in rows:
        item = dict(row)
        try:
            item["content_json"] = json.loads(item.get("content_json") or "{}")
        except Exception:
            pass
        out.append(item)
    return jsonify(out)


@app.route("/api/wedding-sites", methods=["POST"])
def upsert_wedding_site():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    title = str(data.get("title") or "Wedding Website").strip()
    theme = str(data.get("theme") or "classic").strip()
    content_json = json.dumps(data.get("content") or {})
    site_id = data.get("id")

    with get_db() as conn:
        if site_id:
            cur = conn.execute(
                "UPDATE wedding_sites SET title=?, theme=?, content_json=?, updated_at=datetime('now') WHERE id=?",
                (title, theme, content_json, site_id),
            )
            if cur.rowcount == 0:
                return _mutation_not_found("wedding_site", site_id)
        else:
            conn.execute(
                "INSERT INTO wedding_sites (wedding_id, title, theme, content_json, status) VALUES (?,?,?,?, 'draft')",
                (wedding_id, title, theme, content_json),
            )
            site_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": site_id})


@app.route("/api/wedding-sites/generate", methods=["POST"])
def generate_wedding_site():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    instruction = str(data.get("prompt") or "Generate a polished wedding website copy package.")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400

    structured = _build_structured_context(wedding_id)
    recent_chat = _build_recent_chat_context(wedding_id)
    default_payload = {
        "title": "Our Wedding",
        "hero_subtitle": "We are excited to celebrate with you.",
        "story": "Our story is still being finalized.",
        "schedule": [],
        "travel": "Travel and stay details will be shared soon.",
        "faq": [],
    }

    if not openai_client:
        return jsonify({"success": True, "content": default_payload, "source": "fallback"})

    prompt = f"""
You are writing copy for a professional wedding website.
Return ONLY a JSON object with keys:
title, hero_subtitle, story, schedule (array of objects with time and label), travel, faq (array of objects with q and a)

Wedding data:
{structured or "No structured data"}

Recent chat/doc context:
{recent_chat or "No recent chat or docs"}

User instruction:
{instruction}
"""
    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": "Return only valid JSON."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1200,
            temperature=0.35,
        )
        raw = (completion.choices[0].message.content or "").strip()
        content = _parse_json_from_model(raw)
        if not isinstance(content, dict):
            content = default_payload
        return jsonify({"success": True, "content": content, "source": "ai"})
    except Exception:
        return jsonify({"success": True, "content": default_payload, "source": "fallback"})


@app.route("/api/wedding-sites/<int:site_id>/publish", methods=["PUT"])
def publish_wedding_site(site_id):
    data = request.get_json() or {}
    requested_slug = (data.get("slug") or "").strip().lower()
    if requested_slug:
        requested_slug = re.sub(r"[^a-z0-9-]+", "-", requested_slug).strip("-")
    with get_db() as conn:
        row = conn.execute("SELECT * FROM wedding_sites WHERE id=?", (site_id,)).fetchone()
        if not row:
            return _mutation_not_found("wedding_site", site_id)
        slug = requested_slug or row["slug"] or f"wedding-{site_id}"
        suffix = 1
        base_slug = slug
        while True:
            taken = conn.execute("SELECT id FROM wedding_sites WHERE slug=? AND id!=?", (slug, site_id)).fetchone()
            if not taken:
                break
            suffix += 1
            slug = f"{base_slug}-{suffix}"
        conn.execute(
            "UPDATE wedding_sites SET slug=?, status='published', published_at=datetime('now'), updated_at=datetime('now') WHERE id=?",
            (slug, site_id),
        )
    return jsonify({"success": True, "slug": slug, "url": f"{APP_BASE_URL}/w/{slug}"})


@app.route("/w/<slug>")
def render_wedding_site(slug):
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM wedding_sites WHERE slug=? AND status='published' ORDER BY updated_at DESC LIMIT 1",
            (slug,),
        ).fetchone()
    if not row:
        return "<h1>Wedding site not found</h1>", 404
    try:
        content = json.loads(row["content_json"] or "{}")
    except Exception:
        content = {}
    title = html_escape(content.get("title") or row["title"] or "Wedding")
    subtitle = html_escape(content.get("hero_subtitle") or "")
    story = html_escape(content.get("story") or "")
    travel = html_escape(content.get("travel") or "")
    schedule = content.get("schedule") or []
    faq = content.get("faq") or []
    schedule_html = "".join(
        f"<li><strong>{html_escape(str(item.get('time', '')))}</strong> — {html_escape(str(item.get('label', '')))}</li>"
        for item in schedule if isinstance(item, dict)
    )
    faq_html = "".join(
        f"<li><strong>{html_escape(str(item.get('q', '')))}</strong><br/>{html_escape(str(item.get('a', '')))}</li>"
        for item in faq if isinstance(item, dict)
    )
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title}</title>
  <style>
    body {{ margin:0; font-family: Inter, Arial, sans-serif; color:#20182e; background:#f7f3fd; }}
    .hero {{ padding: 56px 20px 34px; text-align:center; background: linear-gradient(145deg,#f8f2ff,#fff6f9); }}
    .hero h1 {{ margin:0; font-size: clamp(2rem, 5vw, 3.2rem); }}
    .hero p {{ margin: 8px 0 0; color:#5f4f78; }}
    .wrap {{ max-width: 920px; margin: 0 auto; padding: 24px 18px 60px; }}
    section {{ background:#fff; border:1px solid #e8dff6; border-radius:14px; padding:20px; margin-top:14px; }}
    h2 {{ margin:0 0 8px; font-size:1.1rem; }}
    ul {{ margin:0; padding-left:18px; line-height:1.6; }}
  </style>
</head>
<body>
  <div class="hero"><h1>{title}</h1><p>{subtitle}</p></div>
  <div class="wrap">
    <section><h2>Our Story</h2><p>{story}</p></section>
    <section><h2>Schedule</h2><ul>{schedule_html or '<li>Schedule coming soon.</li>'}</ul></section>
    <section><h2>Travel & Stay</h2><p>{travel}</p></section>
    <section><h2>FAQ</h2><ul>{faq_html or '<li>FAQ coming soon.</li>'}</ul></section>
  </div>
</body>
</html>"""
    return Response(html, mimetype="text/html")


@app.route("/api/visual-plan-items")
def get_visual_plan_items():
    wedding_id = request.args.get("wedding_id")
    item_type = request.args.get("item_type", "")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        if item_type:
            rows = conn.execute(
                "SELECT * FROM visual_plan_items WHERE wedding_id=? AND item_type=? ORDER BY id DESC",
                (wedding_id, item_type),
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM visual_plan_items WHERE wedding_id=? ORDER BY id DESC",
                (wedding_id,),
            ).fetchall()
    out = []
    for r in rows:
        item = dict(r)
        for field in ["assigned_guests_json", "details_json"]:
            try:
                item[field] = json.loads(item.get(field) or "[]")
            except Exception:
                pass
        out.append(item)
    return jsonify(out)


@app.route("/api/visual-plan-items", methods=["POST"])
def create_visual_plan_item():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    item_type = (data.get("item_type") or "").strip()
    label = (data.get("label") or "").strip()
    if not wedding_id or item_type not in {"seat", "stay"} or not label:
        return jsonify({"error": "wedding_id, item_type(seat|stay), label required"}), 400
    with get_db() as conn:
        conn.execute(
            """
            INSERT INTO visual_plan_items
            (wedding_id, item_type, label, lat, lng, table_number, capacity, assigned_guests_json, details_json, updated_at)
            VALUES (?,?,?,?,?,?,?,?,?,datetime('now'))
            """,
            (
                wedding_id,
                item_type,
                label,
                data.get("lat"),
                data.get("lng"),
                data.get("table_number"),
                data.get("capacity"),
                json.dumps(data.get("assigned_guests") or []),
                json.dumps(data.get("details") or {}),
            ),
        )
        item_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": item_id})


@app.route("/api/visual-plan-items/<int:item_id>", methods=["PUT"])
def update_visual_plan_item(item_id):
    data = request.get_json() or {}
    fields, vals = [], []
    mapping = {
        "label": "label",
        "lat": "lat",
        "lng": "lng",
        "table_number": "table_number",
        "capacity": "capacity",
    }
    for src, dbf in mapping.items():
        if src in data:
            fields.append(f"{dbf}=?")
            vals.append(data[src])
    if "assigned_guests" in data:
        fields.append("assigned_guests_json=?")
        vals.append(json.dumps(data.get("assigned_guests") or []))
    if "details" in data:
        fields.append("details_json=?")
        vals.append(json.dumps(data.get("details") or {}))
    fields.append("updated_at=datetime('now')")
    vals.append(item_id)
    with get_db() as conn:
        cur = conn.execute(f"UPDATE visual_plan_items SET {','.join(fields)} WHERE id=?", vals)
    if cur.rowcount == 0:
        return _mutation_not_found("visual_plan_item", item_id)
    return jsonify({"success": True, "updated": cur.rowcount})


@app.route("/api/visual-plan-items/<int:item_id>", methods=["DELETE"])
def delete_visual_plan_item(item_id):
    with get_db() as conn:
        cur = conn.execute("DELETE FROM visual_plan_items WHERE id=?", (item_id,))
    if cur.rowcount == 0:
        return _mutation_not_found("visual_plan_item", item_id)
    return jsonify({"success": True, "deleted": cur.rowcount})


@app.route("/api/cards")
def list_cards():
    wedding_id = request.args.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM invitation_cards WHERE wedding_id=? ORDER BY created_at DESC",
            (wedding_id,),
        ).fetchall()
    out = []
    for r in rows:
        item = dict(r)
        try:
            item["content_json"] = json.loads(item.get("content_json") or "{}")
        except Exception:
            pass
        out.append(item)
    return jsonify(out)


@app.route("/api/cards/generate", methods=["POST"])
def generate_card_content():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    prompt = str(data.get("prompt") or "").strip()
    preferred_theme = str(data.get("theme") or "classic").strip().lower()
    preferred_title = str(data.get("title") or "").strip()

    default_content = {
        "title": preferred_title or "Wedding Invitation",
        "subtitle": "Together with our families",
        "body": "Join us as we celebrate our wedding day.",
        "theme": preferred_theme if preferred_theme in {"classic", "modern", "royal", "minimal"} else "classic",
        "palette": ["#fef6f9", "#f2dded"],
    }

    if not openai_client:
        return jsonify({"success": True, "content": default_content, "source": "fallback"})

    structured = _build_structured_context(wedding_id) if wedding_id else ""
    recent_chat = _build_recent_chat_context(wedding_id) if wedding_id else ""
    prompt_text = f"""
Create polished invitation card copy for a wedding planning app.
Return ONLY JSON with keys:
title, subtitle, body, theme (classic|modern|royal|minimal), palette (array of 2 hex colors).

User prompt:
{prompt or "Create an elegant wedding invitation."}

Preferred title:
{preferred_title or "Not specified"}

Preferred theme:
{preferred_theme}

Structured context:
{structured or "No structured data"}

Recent chat/doc context:
{recent_chat or "No recent context"}
"""
    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": "Return only valid JSON."},
                {"role": "user", "content": prompt_text},
            ],
            max_tokens=600,
            temperature=0.4,
        )
        raw = (completion.choices[0].message.content or "").strip()
        content = _parse_json_from_model(raw)
        if not isinstance(content, dict):
            content = default_content
        theme = str(content.get("theme") or default_content["theme"]).lower()
        if theme not in {"classic", "modern", "royal", "minimal"}:
            theme = default_content["theme"]
        palette = content.get("palette")
        if not isinstance(palette, list) or len(palette) < 2:
            palette = default_content["palette"]
        content["theme"] = theme
        content["palette"] = [str(palette[0]), str(palette[1])]
        content["title"] = str(content.get("title") or default_content["title"]).strip()
        content["subtitle"] = str(content.get("subtitle") or default_content["subtitle"]).strip()
        content["body"] = str(content.get("body") or default_content["body"]).strip()
        return jsonify({"success": True, "content": content, "source": "ai"})
    except Exception:
        return jsonify({"success": True, "content": default_content, "source": "fallback"})


@app.route("/api/cards", methods=["POST"])
def create_card():
    data = request.get_json() or {}
    wedding_id = (data.get("wedding_id") or "").strip()
    title = (data.get("title") or "").strip()
    if not wedding_id or not title:
        return jsonify({"error": "wedding_id and title required"}), 400
    with get_db() as conn:
        conn.execute(
            """
            INSERT INTO invitation_cards (wedding_id, title, prompt, theme, content_json, created_by)
            VALUES (?,?,?,?,?,?)
            """,
            (
                wedding_id,
                title,
                str(data.get("prompt") or ""),
                str(data.get("theme") or "classic"),
                json.dumps(data.get("content") or {}),
                str(data.get("created_by") or ""),
            ),
        )
        card_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    return jsonify({"success": True, "id": card_id})


@app.route("/api/ai/task-candidates", methods=["POST"])
def ai_task_candidates():
    data = request.get_json() or {}
    text = (data.get("text") or "").strip()
    wedding_id = (data.get("wedding_id") or "").strip()
    if not text:
        return jsonify({"candidates": []})

    def infer_due_date_from_text(raw_text: str) -> str:
        due_match = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", raw_text)
        if due_match:
            return due_match.group(1)
        lower = raw_text.lower()
        today = datetime.now().date()
        if "tomorrow" in lower:
            return (today + timedelta(days=1)).isoformat()
        if "next week" in lower:
            return (today + timedelta(days=7)).isoformat()
        if "this week" in lower:
            return (today + timedelta(days=3)).isoformat()
        weekday_map = {
            "monday": 0, "tuesday": 1, "wednesday": 2, "thursday": 3,
            "friday": 4, "saturday": 5, "sunday": 6,
        }
        current = today.weekday()
        for name, target in weekday_map.items():
            if re.search(rf"\b{name}\b", lower):
                delta = (target - current) % 7
                if delta == 0:
                    delta = 7
                return (today + timedelta(days=delta)).isoformat()
        return ""

    fallback = []
    due_guess = infer_due_date_from_text(text)
    has_general_action = re.search(r"\b(book|finali[sz]e|confirm|send|review|arrange|schedule|create)\b", text, flags=re.I)
    has_meeting_intent = re.search(r"\b(meet|meeting|call|sync|catch[- ]?up|go through|walk through)\b", text, flags=re.I)
    asks_for_time = re.search(r"\b(when|what time|good time|available|tomorrow|next week|this week|monday|tuesday|wednesday|thursday|friday|saturday|sunday)\b", text, flags=re.I)

    if has_meeting_intent and asks_for_time:
        fallback.append({
            "title": "Schedule vendor meeting",
            "description": text,
            "due_date": due_guess,
            "priority": "high",
            "assigned_to": "",
            "category": "Vendors",
            "confidence": 0.82,
        })

    if has_general_action:
        fallback.append({
            "title": text[:80],
            "description": text,
            "due_date": due_guess,
            "priority": "medium",
            "assigned_to": "",
            "category": "Planning",
            "confidence": 0.55,
        })

    if not openai_client:
        return jsonify({"candidates": fallback})

    structured = _build_structured_context(wedding_id) if wedding_id else ""
    prompt = f"""
Extract 0-3 actionable task candidates from this chat message.
Return ONLY a JSON array. Each object:
{{"title":"...","description":"...","due_date":"YYYY-MM-DD or ''","priority":"low|medium|high|urgent","assigned_to":"...","category":"...","confidence":0-1}}

Message:
{text}

Wedding context:
{structured or "No structured data"}

Rules:
- Always detect scheduling intent as a task when the message asks to meet/call/sync or asks "what time"/"good time to meet".
- Prefer concise actionable task titles.
"""
    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": "Return only valid JSON."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=500,
            temperature=0.1,
        )
        parsed = _parse_json_from_model((completion.choices[0].message.content or "").strip())
        if isinstance(parsed, list):
            cands = [c for c in parsed if isinstance(c, dict)]
        elif isinstance(parsed, dict):
            cands = [parsed]
        else:
            cands = []
        return jsonify({"candidates": cands[:3] or fallback})
    except Exception:
        return jsonify({"candidates": fallback})

# ══════════════════════════════════════════════════════════════════════════════
# AI QUERY ENDPOINTS (original + enhanced with structured data)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/ai-query", methods=["POST"])
def ai_query():
    if not openai_client:
        return jsonify({"error": "OPENAI_API_KEY not configured in backend/.env"}), 503

    data      = request.get_json() or {}
    cohort_id = data.get("cohort_id")
    question  = (data.get("question") or "").strip()
    user_id   = data.get("user_id")
    wedding_id = data.get("wedding_id", "")

    if not question:
        return jsonify({"error": "question is required"}), 400

    q_embedding = embed([question])[0]

    context_text, sources = "", []
    try:
        collection = get_or_create_collection(cohort_id)
        results    = collection.query(query_embeddings=[q_embedding], n_results=3)
        docs       = results["documents"][0]
        metas      = results["metadatas"][0]
        if docs:
            context_text = "\n\n---\n\n".join(docs)
            sources      = list({m["filename"] for m in metas})
    except Exception:
        pass

    system_prompt = _build_system_prompt()
    structured_context = _build_structured_context(wedding_id) if wedding_id else ""
    user_prompt = ""
    if structured_context:
        user_prompt += f"Structured wedding data:\n{structured_context}\n\n---\n\n"
    if context_text:
        user_prompt += f"Context from uploaded documents:\n{context_text}\n\n---\n\n"
    user_prompt += f"Question: {question}"

    completion = openai_client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        max_tokens=600,
        temperature=0.4,
    )
    answer = completion.choices[0].message.content

    with get_db() as conn:
        conn.execute(
            "INSERT INTO messages (cohort_id, user_id, content, sender_type) VALUES (?,?,?,?)",
            (cohort_id, None, answer, "ai")
        )

    return jsonify({"response": answer, "sources": sources})


def _build_structured_context(wedding_id):
    """Build structured data context for AI from database."""
    if not wedding_id:
        return ""
    context_parts = []
    try:
        with get_db() as conn:
            # Budget data
            budgets = conn.execute("SELECT * FROM budgets WHERE wedding_id=?", (wedding_id,)).fetchall()
            if budgets:
                budget_lines = ["BUDGET OVERVIEW:"]
                total_alloc, total_spent = 0, 0
                for b in budgets:
                    budget_lines.append(f"  - {b['category']}: allocated ${b['allocated_amount']:.0f}, spent ${b['spent_amount']:.0f}, remaining ${b['allocated_amount']-b['spent_amount']:.0f}")
                    total_alloc += b['allocated_amount']
                    total_spent += b['spent_amount']
                budget_lines.append(f"  TOTAL: allocated ${total_alloc:.0f}, spent ${total_spent:.0f}, remaining ${total_alloc-total_spent:.0f}")
                context_parts.append("\n".join(budget_lines))

            # Tasks data
            tasks = conn.execute("SELECT * FROM wedding_tasks WHERE wedding_id=?", (wedding_id,)).fetchall()
            if tasks:
                task_lines = ["TASKS/TIMELINE:"]
                for t in tasks:
                    task_lines.append(f"  - [{t['status'].upper()}] {t['title']} (due: {t['due_date'] or 'no date'}, assigned: {t['assigned_to'] or 'unassigned'}, priority: {t['priority']})")
                context_parts.append("\n".join(task_lines))

            # Events data
            events = conn.execute("SELECT * FROM wedding_events WHERE wedding_id=? ORDER BY start_time", (wedding_id,)).fetchall()
            if events:
                event_lines = ["WEDDING DAY EVENTS:"]
                for e in events:
                    event_lines.append(f"  - {e['name']}: {e['start_time']} to {e['end_time']} (vendor: {e['vendor_name'] or 'N/A'}, location: {e['location'] or 'TBD'})")
                context_parts.append("\n".join(event_lines))

            # Guest data summary
            guest_count = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=?", (wedding_id,)).fetchone()[0]
            if guest_count:
                attending = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=? AND rsvp_status='attending'", (wedding_id,)).fetchone()[0]
                pending = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=? AND rsvp_status='pending'", (wedding_id,)).fetchone()[0]
                declined = conn.execute("SELECT COUNT(*) FROM guests WHERE wedding_id=? AND rsvp_status='declined'", (wedding_id,)).fetchone()[0]
                meals = conn.execute("SELECT meal_preference, COUNT(*) as cnt FROM guests WHERE wedding_id=? AND rsvp_status='attending' AND meal_preference != '' GROUP BY meal_preference", (wedding_id,)).fetchall()
                guest_lines = [f"GUEST LIST: {guest_count} total, {attending} attending, {pending} pending, {declined} declined"]
                if meals:
                    guest_lines.append("  Meal preferences: " + ", ".join(f"{m['meal_preference']}: {m['cnt']}" for m in meals))
                context_parts.append("\n".join(guest_lines))
    except Exception:
        pass

    return "\n\n".join(context_parts)


def _build_recent_chat_context(wedding_id, limit=25):
    """Collect recent chat and document signals for content-generation tasks."""
    if not wedding_id:
        return ""
    try:
        with get_db() as conn:
            cohort = conn.execute(
                "SELECT id FROM cohorts WHERE lower(name)=lower(?) ORDER BY id DESC LIMIT 1",
                (wedding_id,),
            ).fetchone()
            if not cohort:
                return ""
            cohort_id = cohort["id"]
            rows = conn.execute(
                """
                SELECT sender_type, content, created_at
                FROM messages
                WHERE cohort_id=?
                ORDER BY id DESC
                LIMIT ?
                """,
                (cohort_id, limit),
            ).fetchall()
            docs = conn.execute(
                "SELECT filename, created_at FROM documents WHERE cohort_id=? ORDER BY id DESC LIMIT 8",
                (cohort_id,),
            ).fetchall()
        lines = []
        if rows:
            lines.append("RECENT CHAT CONTEXT:")
            for r in reversed(rows):
                text = re.sub(r"\s+", " ", str(r["content"] or "")).strip()
                if not text:
                    continue
                role = "AI" if str(r["sender_type"]) == "ai" else "User"
                lines.append(f"  - {role}: {text[:240]}")
        if docs:
            lines.append("UPLOADED DOCUMENTS:")
            for d in docs:
                lines.append(f"  - {d['filename']}")
        return "\n".join(lines)
    except Exception:
        return ""


def _build_system_prompt():
    return """You are an expert wedding planning assistant powered by a curated knowledge base of wedding venues, vendors, timelines, budgets, etiquette guides, and real couple experiences.

You have access to STRUCTURED WEDDING DATA including budget allocations, expenses, tasks/timeline, event schedules, and guest lists. Always ground your answers in this real data when available.

**How to answer:**
1. When structured data is provided, prioritize it for factual claims about budgets, tasks, timelines, and guests.
2. For document-based questions, reference the uploaded context. Add inline citations [Source: <document_name>].
3. If multiple sources agree, cite all of them.
4. If information is missing, clearly state what you found and what's missing.
5. Never fabricate vendor names, prices, availability, or policies.

**Budget questions:**
- Always reference actual allocated and spent amounts from the structured data.
- Calculate remaining budget, percentage spent, and flag any overages.
- Compare across categories when asked.

**Timeline/Task questions:**
- Reference actual task statuses, due dates, and assignments.
- Flag overdue tasks proactively.
- Suggest next steps based on pending tasks.

**Event scheduling:**
- Flag any detected time overlaps between events.
- Consider vendor arrival times and setup requirements.

**Guest list questions:**
- Reference actual RSVP counts and meal preferences.
- Calculate totals including plus-ones.

**Response format:**
- Lead with a direct answer, then supporting details.
- Use markdown formatting: **bold** for emphasis, bullet lists with `-`, numbered lists.
- End every response with a "Sources Used" section listing each document referenced.
- If the request is ambiguous and a concrete answer would be unreliable, ask one concise clarifying question first.

**Domain rules:**
- Budget figures must always include the date they were recorded, since pricing changes seasonally.
- Venue availability is time-sensitive — remind users to verify directly with vendors.
- For etiquette questions, note if advice varies by culture or region.
- Flag any potential conflicts proactively.
"""


def _strip_json_fences(raw_text: str) -> str:
    cleaned = re.sub(r"^```(?:json)?\s*", "", raw_text.strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def _parse_json_from_model(raw_text: str):
    cleaned = _strip_json_fences(raw_text)
    try:
        return json.loads(cleaned)
    except Exception:
        pass

    start_obj = cleaned.find("{")
    end_obj = cleaned.rfind("}")
    if start_obj != -1 and end_obj != -1 and end_obj > start_obj:
        snippet = cleaned[start_obj:end_obj + 1]
        try:
            return json.loads(snippet)
        except Exception:
            pass

    start_arr = cleaned.find("[")
    end_arr = cleaned.rfind("]")
    if start_arr != -1 and end_arr != -1 and end_arr > start_arr:
        snippet = cleaned[start_arr:end_arr + 1]
        try:
            return json.loads(snippet)
        except Exception:
            pass

    raise ValueError("Model output did not contain valid JSON.")


# ── API: Streaming AI query (SSE) — enhanced with structured data ───────────
@app.route("/api/ai-query-stream", methods=["POST"])
def ai_query_stream():
    if not openai_client:
        return jsonify({"error": "OPENAI_API_KEY not configured in backend/.env"}), 503

    data      = request.get_json() or {}
    cohort_id = data.get("cohort_id")
    question  = (data.get("question") or "").strip()
    user_id   = data.get("user_id")
    wedding_id = data.get("wedding_id", "")

    if not question:
        return jsonify({"error": "question is required"}), 400

    q_embedding = embed([question])[0]

    context_text, sources = "", []
    try:
        collection = get_or_create_collection(cohort_id)
        results    = collection.query(query_embeddings=[q_embedding], n_results=3)
        docs       = results["documents"][0]
        metas      = results["metadatas"][0]
        if docs:
            context_text = "\n\n---\n\n".join(docs)
            sources      = list({m["filename"] for m in metas})
    except Exception:
        pass

    system_prompt = _build_system_prompt()
    structured_context = _build_structured_context(wedding_id) if wedding_id else ""

    user_prompt = ""
    if structured_context:
        user_prompt += f"Structured wedding data:\n{structured_context}\n\n---\n\n"
    if context_text:
        user_prompt += f"Context from uploaded documents:\n{context_text}\n\n---\n\n"
    user_prompt += f"Question: {question}"

    def generate():
        full_answer = ""
        yield f"data: {json.dumps({'type': 'sources', 'sources': sources})}\n\n"

        stream = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            max_tokens=600,
            temperature=0.4,
            stream=True,
        )

        for chunk in stream:
            delta = chunk.choices[0].delta
            if delta.content:
                full_answer += delta.content
                yield f"data: {json.dumps({'type': 'chunk', 'content': delta.content})}\n\n"

        with get_db() as conn:
            conn.execute(
                "INSERT INTO messages (cohort_id, user_id, content, sender_type) VALUES (?,?,?,?)",
                (cohort_id, None, full_answer, "ai")
            )

        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── AI Coordination: smart suggestions endpoint ──────────────────────────────
@app.route("/api/ai-suggestions", methods=["POST"])
def ai_suggestions():
    """Generate proactive AI suggestions based on current wedding state."""
    if not openai_client:
        return jsonify({"error": "OPENAI_API_KEY not configured"}), 503

    data = request.get_json() or {}
    wedding_id = data.get("wedding_id")
    if not wedding_id:
        return jsonify({"error": "wedding_id required"}), 400

    structured = _build_structured_context(wedding_id)
    if not structured:
        return jsonify({"suggestions": [], "message": "No data yet. Add tasks, events, and budget items first."})

    prompt = f"""Analyze the following wedding planning data and provide 3-5 actionable suggestions. Focus on:
1. Schedule conflicts or tight timings
2. Budget warnings (overages, categories with no allocation)
3. Overdue or upcoming tasks
4. Missing vendors or services
5. Guest list issues

Wedding data:
{structured}

Return ONLY a JSON array of suggestion objects. Each object: {{"type": "warning"|"tip"|"reminder", "title": "short title", "message": "detailed suggestion"}}
"""

    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": "You are a proactive wedding planning assistant. Analyze data and give actionable suggestions. Return ONLY valid JSON."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=800,
            temperature=0.3,
        )
        raw = completion.choices[0].message.content.strip()
        cleaned = re.sub(r"^```(?:json)?\s*", "", raw)
        cleaned = re.sub(r"\s*```$", "", cleaned)
        suggestions = json.loads(cleaned)
        if not isinstance(suggestions, list):
            suggestions = [suggestions]
    except Exception as e:
        suggestions = [{"type": "tip", "title": "Keep planning!", "message": f"Add more data for smarter suggestions. ({str(e)[:50]})"}]

    return jsonify({"suggestions": suggestions})


@app.route("/api/ai-copilot-plan", methods=["POST"])
def ai_copilot_plan():
    data = request.get_json() or {}
    instruction = (data.get("instruction") or "").strip()
    wedding_id = data.get("wedding_id", "")
    suggestion_context = data.get("suggestion_context") or {}

    if not instruction:
        return jsonify({"error": "instruction is required"}), 400

    lowered = instruction.lower()
    if ("generate" in lowered or "create" in lowered or "make" in lowered) and "doc" in lowered:
        if "pdf" not in lowered and "docx" not in lowered and "word" not in lowered:
            return jsonify({
                "status": "needs_clarification",
                "summary": "Document format is missing.",
                "question": "Do you want the generated document as a PDF or a DOCX file?",
                "actions": [],
            })

    if not openai_client:
        return jsonify({"status": "not_actionable", "summary": "AI backend is unavailable.", "actions": []})

    structured_context = _build_structured_context(wedding_id) if wedding_id else ""
    suggestion_blob = json.dumps(suggestion_context)

    system_prompt = """You are the execution planner for a wedding planning copilot.
Return ONLY one JSON object with this exact shape:
{
  "status": "ready" | "needs_clarification" | "not_actionable",
  "summary": "short summary",
  "question": "only when status is needs_clarification",
  "actions": [ { "type": "..." } ]
}

Allowed action types:
- remove_recent_expenses: {"type":"remove_recent_expenses","budget_categories":[...],"count":number,"until_within_budget":true|false}
- notify_group_chat: {"type":"notify_group_chat","message":"...","group_name":"optional","vendor_names":[...]}
- notify_vendor_chat: {"type":"notify_vendor_chat","message":"...","vendor_names":[...]}
- send_guest_invites: {"type":"send_guest_invites","guest_ids":[...],"send_email":true|false}
- post_announcement: {"type":"post_announcement","message":"...","metadata":{...}}
- create_task: {"type":"create_task","title":"...","description":"...","due_date":"YYYY-MM-DD","priority":"low|medium|high|urgent","assigned_to":"...","category":"..."}
- update_task: {"type":"update_task","task_id":number,"status":"pending|in_progress|completed|overdue","due_date":"YYYY-MM-DD","assigned_to":"..."}
- extend_overdue_tasks: {"type":"extend_overdue_tasks","days":number,"task_ids":[...]}
- reallocate_budget: {"type":"reallocate_budget","from_category":"...","to_category":"...","amount":number}
- switch_tab: {"type":"switch_tab","tab":"chat|budget|tasks|guests|discover|website|plans|cards|announcements"}
- generate_document: {"type":"generate_document","prompt":"...","output_format":"pdf|docx"}

Rules:
- Use not_actionable for informational questions.
- Use needs_clarification when a critical detail is missing.
- For direct instructions, return concrete actions.
- Do not emit unsupported action types.
"""

    user_prompt = f"""User instruction:
{instruction}

Structured wedding data:
{structured_context or "No structured data available."}

Current copilot suggestion context:
{suggestion_blob}
"""

    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=900,
            temperature=0.1,
        )
        raw = (completion.choices[0].message.content or "").strip()
        parsed = _parse_json_from_model(raw)
        if not isinstance(parsed, dict):
            raise ValueError("Expected JSON object")

        status = parsed.get("status", "not_actionable")
        if status not in {"ready", "needs_clarification", "not_actionable"}:
            status = "not_actionable"

        actions = parsed.get("actions", [])
        if not isinstance(actions, list):
            actions = []
        actions = [a for a in actions if isinstance(a, dict) and a.get("type")]

        payload = {
            "status": status,
            "summary": str(parsed.get("summary", "")).strip(),
            "question": str(parsed.get("question", "")).strip(),
            "actions": actions,
        }
        if payload["status"] == "ready" and not actions:
            payload["status"] = "not_actionable"
        return jsonify(payload)
    except Exception as e:
        return jsonify({
            "status": "not_actionable",
            "summary": f"Planner fallback used: {str(e)[:120]}",
            "actions": [],
        })


# ── Keyword helpers for doc generation ────────────────────────────────────────
STOP_WORDS = {
    "a","an","and","are","as","at","be","by","for","from","has","have","in",
    "is","it","its","of","on","or","that","the","to","was","were","will",
    "with","what","when","where","who","how","why","we","you","your","our",
    "this","these","those","they","their","i","me","my","about","create",
    "generate","doc","document","make","write","please","can","could","would",
    "should","want","need","like","into","also","some","all","do","does","did",
}


def extract_keywords(text: str) -> list:
    tokens = re.findall(r"[a-z0-9]+", text.lower())
    return [t for t in tokens if len(t) > 2 and t not in STOP_WORDS]


def search_messages_by_keywords(cohort_id, keywords, limit=30):
    if not keywords:
        return []
    clauses = " OR ".join(["content LIKE ?"] * len(keywords))
    params  = [f"%{kw}%" for kw in keywords]
    with get_db() as conn:
        rows = conn.execute(f"""
            SELECT content, sender_type, created_at
            FROM   messages
            WHERE  cohort_id = ? AND ({clauses})
            ORDER  BY created_at DESC
            LIMIT  ?
        """, [cohort_id] + params + [limit]).fetchall()
    return [dict(r) for r in rows]


def build_docx(title: str, sections: list, filename: str) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1D, 0x1A, 0x16)
    style.paragraph_format.space_after = Pt(6)

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.runs[0]
    run.font.color.rgb = RGBColor(0xB3, 0x4D, 0x2E)

    ts = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts.runs[0].font.size  = Pt(9)
    ts.runs[0].font.color.rgb = RGBColor(0x6B, 0x62, 0x55)

    doc.add_paragraph("")

    for section in sections:
        heading_text = section.get("heading", "")
        body_text    = section.get("body", "")
        if heading_text:
            h = doc.add_heading(heading_text, level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x8E, 0x37, 0x1C)
        if body_text:
            for para_text in body_text.split("\n"):
                stripped = para_text.strip()
                if not stripped:
                    continue
                if stripped.startswith("- ") or stripped.startswith("• "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif re.match(r"^\d+[\.\)]\s", stripped):
                    doc.add_paragraph(re.sub(r"^\d+[\.\)]\s*", "", stripped), style="List Number")
                else:
                    doc.add_paragraph(stripped)

    filepath = DOCS_OUT_DIR / filename
    doc.save(str(filepath))
    return filepath


def _wrap_plain_line(text: str, max_len: int = 95) -> list:
    words = text.split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        if len(current) + 1 + len(word) <= max_len:
            current += " " + word
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _sections_to_lines(title: str, sections: list) -> list:
    lines = [
        title,
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        "",
    ]
    for section in sections:
        heading_text = str(section.get("heading", "")).strip()
        body_text = str(section.get("body", "")).strip()
        if heading_text:
            lines.append(heading_text.upper())
        if body_text:
            for para_text in body_text.split("\n"):
                stripped = para_text.strip()
                if not stripped:
                    lines.append("")
                elif stripped.startswith("- "):
                    lines.extend(_wrap_plain_line(f"• {stripped[2:]}", max_len=92))
                else:
                    lines.extend(_wrap_plain_line(stripped))
        lines.append("")
    return lines


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_pdf(title: str, sections: list, filename: str) -> Path:
    lines = _sections_to_lines(title, sections)
    lines_per_page = 48
    pages = [lines[i:i + lines_per_page] for i in range(0, len(lines), lines_per_page)] or [[""]]

    objects = {}
    page_count = len(pages)
    font_obj_num = 3 + page_count * 2
    max_obj_num = font_obj_num

    kids = " ".join(f"{3 + idx * 2} 0 R" for idx in range(page_count))
    objects[1] = "<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>"

    for idx, page_lines in enumerate(pages):
        page_obj_num = 3 + idx * 2
        content_obj_num = page_obj_num + 1

        stream_lines = [
            "BT",
            "/F1 11 Tf",
            "14 TL",
            "72 760 Td",
        ]
        for line in page_lines:
            stream_lines.append(f"({_pdf_escape(line)}) Tj")
            stream_lines.append("T*")
        stream_lines.append("ET")
        stream_text = "\n".join(stream_lines)
        stream_bytes = stream_text.encode("latin-1", errors="replace")

        objects[page_obj_num] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_obj_num} 0 R >> >> "
            f"/Contents {content_obj_num} 0 R >>"
        )
        objects[content_obj_num] = (
            f"<< /Length {len(stream_bytes)} >>\nstream\n"
            + stream_text
            + "\nendstream"
        )

    objects[font_obj_num] = "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    pdf = bytearray()
    pdf.extend(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = {0: 0}

    for obj_num in range(1, max_obj_num + 1):
        offsets[obj_num] = len(pdf)
        pdf.extend(f"{obj_num} 0 obj\n".encode("latin-1"))
        pdf.extend(str(objects[obj_num]).encode("latin-1", errors="replace"))
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {max_obj_num + 1}\n".encode("latin-1"))
    pdf.extend(b"0000000000 65535 f \n")
    for obj_num in range(1, max_obj_num + 1):
        pdf.extend(f"{offsets[obj_num]:010d} 00000 n \n".encode("latin-1"))
    pdf.extend(
        (
            f"trailer\n<< /Size {max_obj_num + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("latin-1")
    )

    filepath = DOCS_OUT_DIR / filename
    filepath.write_bytes(pdf)
    return filepath


@app.route("/api/generate-doc", methods=["POST"])
def generate_doc():
    if not openai_client:
        return jsonify({"error": "OPENAI_API_KEY not configured in backend/.env"}), 503

    data      = request.get_json() or {}
    cohort_id = data.get("cohort_id")
    user_id   = data.get("user_id")
    prompt    = (data.get("prompt") or "").strip()
    wedding_id = data.get("wedding_id", "")
    output_format = (data.get("output_format") or "").strip().lower()

    if not prompt:
        return jsonify({"error": "prompt is required"}), 400
    if not cohort_id:
        return jsonify({"error": "cohort_id is required"}), 400

    if not output_format:
        lowered_prompt = prompt.lower()
        if "pdf" in lowered_prompt:
            output_format = "pdf"
        elif "docx" in lowered_prompt or "word" in lowered_prompt:
            output_format = "docx"
    if output_format not in {"docx", "pdf"}:
        return jsonify({
            "error": "output_format must be 'docx' or 'pdf'",
            "needs_clarification": True,
            "question": "Do you want a PDF or a DOCX file?",
        }), 400

    keywords = extract_keywords(prompt)
    matched_msgs = search_messages_by_keywords(cohort_id, keywords, limit=30)
    chat_context = ""
    if matched_msgs:
        chat_lines = []
        for m in matched_msgs:
            tag = m["sender_type"].upper()
            chat_lines.append(f"[{tag} – {m['created_at']}] {m['content']}")
        chat_context = "\n".join(chat_lines)

    kb_context = ""
    kb_sources = []
    try:
        q_embedding = embed([prompt])[0]
        collection  = get_or_create_collection(cohort_id)
        results     = collection.query(query_embeddings=[q_embedding], n_results=5)
        docs  = results["documents"][0]
        metas = results["metadatas"][0]
        if docs:
            kb_context = "\n\n---\n\n".join(docs)
            kb_sources = list({m["filename"] for m in metas})
    except Exception:
        pass

    system_prompt = """You are a professional document writer for a wedding planning team.
You will be given context from multiple sources:
  1) CHAT HISTORY — messages exchanged inside the wedding planning workspace.
  2) KNOWLEDGE BASE — excerpts from uploaded planning documents.
  3) STRUCTURED DATA — budget, tasks, events, and guest information.

Your job is to generate a polished, well-structured document based on the user's request.

**Output format — strict JSON array:**
Return ONLY a JSON array of section objects. No markdown fences, no commentary outside the JSON.
Each object has two keys:
  - "heading": a short section title (string)
  - "body": the section content (string, may include newlines, bullet lines starting with "- ")

**Rules:**
- Use ONLY information present in the provided context. Do not fabricate names, prices, or dates.
- If information is missing, note it clearly (e.g., "To be confirmed").
- Organise content logically with clear headings.
- Write in a professional yet warm tone appropriate for wedding planning.
- Be thorough — include all relevant details from the context.
"""

    combined_context = ""
    structured = _build_structured_context(wedding_id) if wedding_id else ""
    if structured:
        combined_context += f"=== STRUCTURED DATA ===\n{structured}\n\n"
    if chat_context:
        combined_context += f"=== CHAT HISTORY ===\n{chat_context}\n\n"
    if kb_context:
        combined_context += f"=== KNOWLEDGE BASE ===\n{kb_context}\n\n"

    if not combined_context.strip():
        return jsonify({"error": "No relevant content found. Try a different description."}), 404

    user_prompt = f"{combined_context}---\n\nUser request: {prompt}"

    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": user_prompt},
            ],
            max_tokens=2000,
            temperature=0.3,
        )
        raw_answer = completion.choices[0].message.content.strip()
    except Exception as e:
        return jsonify({"error": f"OpenAI API error: {str(e)}"}), 502

    try:
        parsed = _parse_json_from_model(raw_answer)
        if isinstance(parsed, list):
            sections = parsed
        elif isinstance(parsed, dict):
            sections = [parsed]
        else:
            raise ValueError("Expected JSON array/object")
    except Exception:
        sections = [{"heading": "Document", "body": raw_answer}]

    title_words = prompt.split()[:8]
    doc_title   = " ".join(title_words).title()
    if len(prompt.split()) > 8:
        doc_title += "..."

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^a-z0-9]+", "_", prompt.lower())[:40].strip("_")
    extension = "pdf" if output_format == "pdf" else "docx"
    filename = f"{safe_name}_{timestamp}.{extension}"
    filepath = build_pdf(doc_title, sections, filename) if output_format == "pdf" else build_docx(doc_title, sections, filename)

    with get_db() as conn:
        conn.execute(
            "INSERT INTO generated_docs (cohort_id, user_id, prompt, filename, filepath) VALUES (?,?,?,?,?)",
            (cohort_id, user_id, prompt, filename, str(filepath))
        )
        doc_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        source_note = f" (Sources: {', '.join(kb_sources)})" if kb_sources else ""
        ai_msg = f"Document generated: **{filename}**{source_note}"
        conn.execute(
            "INSERT INTO messages (cohort_id, user_id, content, sender_type) VALUES (?,?,?,?)",
            (cohort_id, None, ai_msg, "ai")
        )

    return jsonify({
        "success":  True,
        "doc_id":   doc_id,
        "filename": filename,
        "output_format": output_format,
        "sources":  kb_sources,
        "sections": len(sections),
    })


@app.route("/api/download-doc/<int:doc_id>")
def download_doc(doc_id):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM generated_docs WHERE id = ?", (doc_id,)).fetchone()
    if not row:
        return jsonify({"error": "Document not found"}), 404
    filepath = Path(row["filepath"])
    if not filepath.exists():
        return jsonify({"error": "File missing from server"}), 404
    return send_file(str(filepath), as_attachment=True, download_name=row["filename"])


@app.route("/api/generated-docs")
def list_generated_docs():
    cohort_id = request.args.get("cohort_id")
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, prompt, filename, created_at FROM generated_docs WHERE cohort_id=? ORDER BY created_at DESC",
            (cohort_id,)
        ).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/knowledge-base")
def knowledge_base():
    cohort_id = request.args.get("cohort_id")
    with get_db() as conn:
        docs = conn.execute(
            "SELECT filename, chunk_count, created_at FROM documents WHERE cohort_id=? ORDER BY created_at DESC",
            (cohort_id,)
        ).fetchall()
    return jsonify([dict(d) for d in docs])


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    init_db()
    key_status = "configured" if OPENAI_API_KEY else "MISSING – add it to backend/.env"
    print(f"\n  WedBoard – AI Wedding Planner Backend")
    print(f"  OpenAI key  : {key_status}")
    print(f"  SQLite db   : {DB_PATH}")
    print(f"  ChromaDB    : {CHROMA_PATH}")
    print(f"  Generated   : {DOCS_OUT_DIR}")
    print(f"  Open        : http://localhost:5000\n")
    app.run(debug=True, port=5000)
